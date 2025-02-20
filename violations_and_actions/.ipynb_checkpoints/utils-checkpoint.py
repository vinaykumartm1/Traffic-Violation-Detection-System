import easyocr
import os
import cv2 
from docx import Document
from datetime import datetime
import pandas as pd 
import numpy as np
from difflib import get_close_matches
from email.message import EmailMessage
import ssl
import smtplib
from email.mime.image import MIMEImage


ROOT_DIR = "./"

def closeMatches(patterns, word):
    matc = get_close_matches(word, patterns)[0]
    return matc
    
     
def make_doc(result_set):
    
    df = pd.read_excel('database.xlsx')
    rns = np.array(df.iloc[:, 0])
    
    for item in result_set:
        fname, result = item[0], item[1]
        try:
            matched = closeMatches(rns, result)
        except:
            continue
        
        match_row = df[df['Rno']==matched]
        
        im_path = './images/'+fname
        per_path = './person/'+fname
        now = datetime.now()

        doc = Document()

        doc.add_heading('Issued by Traffic Regulations Authority', 1)

        table = doc.add_table(rows=1, cols=2)

        data = [
            ['Registration Number', str(matched)],
            ['Name of Owner', match_row['Name']],
            ['Address', match_row['Address']],
            # ['Vehicle Class', 'Bike/Scooter'],
            # ['Date and Time', now.strftime("%d/%m/%Y %H:%M:%S")],
            ['Riding a motor cycle without helmet', '1000'],
            ['Total Fine Amount', '1000 INR']  
        ]
        
        row = table.rows[0].cells
        row[0].text = 'Specification'
        row[1].text = 'Value'
        
        for ref, val in data:
            row = table.add_row().cells
            row[0].text = ref 
            row[1].text = val
            
        
        doc.add_picture(im_path)
        doc.add_picture(per_path)
        save_dir = './challan/{0}.docx'.format(fname)
        doc.save(save_dir)
        send_email(save_dir, match_row['Email'])

    return True
    

def perform_ocr(): 
    im_dir = ROOT_DIR+'images/' 
    reader = easyocr.Reader(['en'])
    fnames = os.listdir(im_dir)
    result_set = []
    for name in fnames:
        print(name)
        im = cv2.imread(im_dir+name)
        result = reader.readtext(im, paragraph="False")
        try:
            result = result[0][-1].replace(" ", '').upper()
            result_set.append([name, result])
            
        except:
            continue
        
    print(result_set)
    return result_set


def send_email(cont_path, mail_receiver):

    mail_sender = ''
    mail_password = ''

    
    subject = 'Violation Detection'
    
    body = """
        A ticket has been issued for your violation of traffic rules
    """

    em = EmailMessage()
    em['From'] = mail_sender
    em['To'] = mail_receiver
    em['subject'] = subject

    with open(cont_path, 'rb') as content_file:
        content = content_file.read()
        em.add_attachment(content, maintype='application', subtype='docx', filename='ticket.docx')

    context = ssl.create_default_context()
    with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
        smtp.login(mail_sender, mail_password)
        smtp.send_message(em)
